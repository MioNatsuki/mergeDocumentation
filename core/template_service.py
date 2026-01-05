"""
Servicio para manejo de plantillas Word (.docx) con placeholders {{nombre}}
"""
import os
import tempfile
import subprocess
import shutil
from docx import Document
import re
from typing import Dict, List, Tuple, Optional
from datetime import datetime
import uuid

class TemplateService:
    """
    Servicio para procesar plantillas Word con datos dinámicos
    """
    
    @staticmethod
    def extraer_placeholders(word_path: str) -> List[str]:
        """
        Extrae todos los {{placeholders}} del documento Word
        Retorna: Lista de placeholders únicos
        """
        try:
            if not os.path.exists(word_path):
                print(f"Archivo no encontrado: {word_path}")
                return []
            
            doc = Document(word_path)
            placeholders = set()  # Usar set para únicos
            
            # Buscar en párrafos
            for paragraph in doc.paragraphs:
                matches = re.findall(r'\{\{(\w+)\}\}', paragraph.text)
                placeholders.update(matches)
            
            # Buscar en tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        matches = re.findall(r'\{\{(\w+)\}\}', cell.text)
                        placeholders.update(matches)
            
            # Buscar en headers (si existen)
            try:
                for section in doc.sections:
                    for paragraph in section.header.paragraphs:
                        matches = re.findall(r'\{\{(\w+)\}\}', paragraph.text)
                        placeholders.update(matches)
                    for paragraph in section.footer.paragraphs:
                        matches = re.findall(r'\{\{(\w+)\}\}', paragraph.text)
                        placeholders.update(matches)
            except:
                pass  # Algunos docs pueden no tener headers/footers
            
            return list(placeholders)
            
        except Exception as e:
            print(f"Error extrayendo placeholders de {word_path}: {e}")
            return []
    
    @staticmethod
    def validar_placeholders(placeholders: List[str], columnas_padron: List[str]) -> Dict:
        """
        Valida placeholders contra columnas del padrón
        Retorna: {"validos": [...], "no_encontrados": [...]}
        """
        columnas_lower = [col.lower() for col in columnas_padron]
        resultados = {"validos": [], "no_encontrados": []}
        
        for placeholder in placeholders:
            placeholder_lower = placeholder.lower()
            
            # Buscar coincidencias
            coincidencias = []
            for i, col in enumerate(columnas_lower):
                if placeholder_lower in col or col in placeholder_lower:
                    coincidencias.append(columnas_padron[i])
            
            if coincidencias:
                resultados["validos"].append({
                    "placeholder": placeholder,
                    "coincidencias": coincidencias,
                    "sugerencia": coincidencias[0]  # Primera coincidencia como sugerencia
                })
            else:
                resultados["no_encontrados"].append(placeholder)
        
        return resultados
    
    @staticmethod
    def generar_docx_con_datos(
        word_path: str, 
        mapeo: Dict,  # {"nombre": "nombre_completo", ...}
        datos: Dict,
        output_path: str = None
    ) -> str:
        """
        Genera .docx con datos reemplazados
        Retorna: Ruta al archivo generado o None en error
        """
        try:
            if not os.path.exists(word_path):
                raise FileNotFoundError(f"Plantilla no encontrada: {word_path}")
            
            # Cargar documento
            doc = Document(word_path)
            
            # Función helper para reemplazar
            def reemplazar_texto(texto: str) -> str:
                """Reemplaza todos los placeholders en un texto"""
                if not texto:
                    return texto
                
                resultado = texto
                for placeholder, columna in mapeo.items():
                    if not columna:
                        continue
                    
                    # Obtener valor del dato
                    valor = datos.get(columna, "")
                    if valor is None:
                        valor = ""
                    
                    # Reemplazar placeholder con valor
                    placeholder_completo = f"{{{{{placeholder}}}}}"
                    if placeholder_completo in resultado:
                        resultado = resultado.replace(placeholder_completo, str(valor))
                
                return resultado
            
            # 1. Reemplazar en párrafos
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    paragraph.text = reemplazar_texto(paragraph.text)
                    
                    # También reemplazar en runs (para mantener formatos)
                    for run in paragraph.runs:
                        if run.text:
                            run.text = reemplazar_texto(run.text)
            
            # 2. Reemplazar en tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            cell.text = reemplazar_texto(cell.text)
                            
                            # Reemplazar en runs de celdas
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if run.text:
                                        run.text = reemplazar_texto(run.text)
            
            # 3. Reemplazar en headers y footers (si existen)
            try:
                for section in doc.sections:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text:
                            paragraph.text = reemplazar_texto(paragraph.text)
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text:
                            paragraph.text = reemplazar_texto(paragraph.text)
            except:
                pass  # Algunos documentos no tienen headers/footers
            
            # 4. Guardar documento
            if not output_path:
                # Crear nombre temporal único
                temp_dir = tempfile.gettempdir()
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"temp_{timestamp}_{os.path.basename(word_path)}"
                output_path = os.path.join(temp_dir, filename)
            
            # Asegurar directorio de salida
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Guardar
            doc.save(output_path)
            print(f"✅ Documento generado: {output_path}")
            return output_path
            
        except Exception as e:
            print(f"❌ Error generando docx: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    @staticmethod
    def convertir_docx_a_pdf(docx_path: str, pdf_path: str = None) -> Tuple[bool, str]:
        """
        Convierte .docx a .pdf usando LibreOffice
        Retorna: (éxito, ruta_pdf o mensaje_error)
        """
        try:
            if not os.path.exists(docx_path):
                return False, f"Archivo no existe: {docx_path}"
            
            # Si no se especifica pdf_path, usar mismo nombre con extensión .pdf
            if not pdf_path:
                pdf_path = docx_path.replace('.docx', '.pdf')
            
            # Verificar si LibreOffice está instalado
            libreoffice_paths = [
                'soffice',  # Linux/macOS
                'libreoffice',  # Alternativa
                'C:\\Program Files\\LibreOffice\\program\\soffice.exe',  # Windows
                'C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe'  # Windows 32-bit
            ]
            
            cmd = None
            for lo_path in libreoffice_paths:
                try:
                    subprocess.run([lo_path, '--version'], 
                                 capture_output=True, check=True, timeout=5)
                    cmd = lo_path
                    break
                except:
                    continue
            
            if not cmd:
                # Fallback: usar python-docx + reportlab (conversión básica)
                return TemplateService._convertir_con_reportlab(docx_path, pdf_path)
            
            # Usar LibreOffice para conversión
            print(f"🔧 Convirtiendo {docx_path} a PDF...")
            
            # Crear comando
            conversion_cmd = [
                cmd, '--headless', '--convert-to', 'pdf',
                '--outdir', os.path.dirname(pdf_path),
                docx_path
            ]
            
            # Ejecutar conversión
            result = subprocess.run(
                conversion_cmd,
                capture_output=True,
                text=True,
                timeout=30  # Timeout de 30 segundos
            )
            
            if result.returncode == 0:
                # LibreOffice crea archivo en directorio de salida
                # con el mismo nombre pero extensión .pdf
                expected_pdf = os.path.join(
                    os.path.dirname(pdf_path),
                    os.path.basename(docx_path).replace('.docx', '.pdf')
                )
                
                if os.path.exists(expected_pdf):
                    # Mover al destino deseado si es diferente
                    if expected_pdf != pdf_path:
                        shutil.move(expected_pdf, pdf_path)
                    
                    print(f"✅ PDF generado: {pdf_path}")
                    return True, pdf_path
                else:
                    return False, f"Archivo PDF no creado en: {expected_pdf}"
            else:
                error_msg = f"Error LibreOffice: {result.stderr}"
                print(f"❌ {error_msg}")
                return False, error_msg
                
        except subprocess.TimeoutExpired:
            return False, "Timeout en conversión PDF"
        except Exception as e:
            error_msg = f"Error convirtiendo a PDF: {str(e)}"
            print(f"❌ {error_msg}")
            return False, error_msg
    
    @staticmethod
    def _convertir_con_reportlab(docx_path: str, pdf_path: str) -> Tuple[bool, str]:
        """Fallback usando ReportLab (conversión básica de texto)"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from docx import Document
            
            print(f"⚠️ Usando fallback ReportLab para {docx_path}")
            
            # Leer contenido del docx
            doc = Document(docx_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    text_content.append(" | ".join(row_text))
            
            # Crear PDF básico
            c = canvas.Canvas(pdf_path, pagesize=letter)
            y = 750  # Posición inicial
            
            for line in text_content:
                if y < 50:  # Nueva página
                    c.showPage()
                    y = 750
                
                c.drawString(50, y, line[:100])  # Limitar longitud
                y -= 15
            
            c.save()
            
            print(f"⚠️ PDF básico generado (sin formato): {pdf_path}")
            return True, pdf_path
            
        except Exception as e:
            return False, f"Error fallback ReportLab: {str(e)}"
    
    @staticmethod
    def generar_pdf_directo(
        word_path: str,
        mapeo: Dict,
        datos: Dict,
        pdf_output_path: str = None
    ) -> Tuple[bool, str, str]:
        """
        Genera PDF directamente: docx → reemplazar datos → pdf
        Retorna: (éxito, ruta_pdf, mensaje)
        """
        temp_docx = None
        try:
            # 1. Generar docx con datos
            print(f"📝 Generando docx con datos...")
            temp_docx = TemplateService.generar_docx_con_datos(word_path, mapeo, datos)
            
            if not temp_docx or not os.path.exists(temp_docx):
                return False, "", "Error generando documento Word con datos"
            
            # 2. Convertir a PDF
            print(f"🔄 Convirtiendo a PDF...")
            if not pdf_output_path:
                # Crear nombre único para PDF
                temp_dir = tempfile.gettempdir()
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                pdf_output_path = os.path.join(temp_dir, f"documento_{timestamp}.pdf")
            
            exito, resultado = TemplateService.convertir_docx_a_pdf(temp_docx, pdf_output_path)
            
            # 3. Limpiar archivo temporal
            if temp_docx and os.path.exists(temp_docx):
                try:
                    os.remove(temp_docx)
                except:
                    pass
            
            if exito:
                return True, resultado, "PDF generado exitosamente"
            else:
                return False, "", f"Error convirtiendo a PDF: {resultado}"
                
        except Exception as e:
            # Limpiar en caso de error
            if temp_docx and os.path.exists(temp_docx):
                try:
                    os.remove(temp_docx)
                except:
                    pass
            
            error_msg = f"Error generando PDF: {str(e)}"
            print(f"❌ {error_msg}")
            return False, "", error_msg
    
    @staticmethod
    def copiar_plantilla_a_destino(word_path: str, destino_dir: str, nombre_plantilla: str = None) -> str:
        """
        Copia archivo Word a directorio de plantillas
        Retorna: Ruta final del archivo
        """
        try:
            if not os.path.exists(word_path):
                raise FileNotFoundError(f"Archivo fuente no existe: {word_path}")
            
            # Crear directorio si no existe
            os.makedirs(destino_dir, exist_ok=True)
            
            # Generar nombre único si no se proporciona
            if not nombre_plantilla:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                unique_id = str(uuid.uuid4())[:8]
                nombre_plantilla = f"plantilla_{timestamp}_{unique_id}.docx"
            elif not nombre_plantilla.lower().endswith('.docx'):
                nombre_plantilla += '.docx'
            
            # Ruta destino
            destino_path = os.path.join(destino_dir, nombre_plantilla)
            
            # Copiar archivo
            shutil.copy2(word_path, destino_path)
            
            print(f"✅ Plantilla copiada a: {destino_path}")
            return destino_path
            
        except Exception as e:
            print(f"❌ Error copiando plantilla: {e}")
            raise
    
    @staticmethod
    def generar_lote_pdfs(
        word_path: str,
        mapeo: Dict,
        lista_datos: List[Dict],
        output_dir: str,
        callback_progreso=None
    ) -> Dict:
        """
        Genera múltiples PDFs (uno por registro)
        Retorna: {
            'total': int,
            'exitosos': int,
            'fallidos': int,
            'errores': List[str],
            'archivos_generados': List[str]
        }
        """
        resultados = {
            'total': len(lista_datos),
            'exitosos': 0,
            'fallidos': 0,
            'errores': [],
            'archivos_generados': []
        }
        
        # Crear directorio de salida
        os.makedirs(output_dir, exist_ok=True)
        
        for i, datos in enumerate(lista_datos):
            try:
                # Obtener cuenta para nombre de archivo
                cuenta = datos.get('cuenta', f'doc_{i+1}')
                nombre_pdf = f"{cuenta}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
                pdf_path = os.path.join(output_dir, nombre_pdf)
                
                # Generar PDF
                exito, ruta_pdf, mensaje = TemplateService.generar_pdf_directo(
                    word_path, mapeo, datos, pdf_path
                )
                
                if exito:
                    resultados['exitosos'] += 1
                    resultados['archivos_generados'].append({
                        'archivo': nombre_pdf,
                        'ruta': ruta_pdf,
                        'cuenta': cuenta
                    })
                else:
                    resultados['fallidos'] += 1
                    resultados['errores'].append(f"{cuenta}: {mensaje}")
                
                # Callback de progreso
                if callback_progreso:
                    callback_progreso(i + 1, len(lista_datos), cuenta, exito)
                    
            except Exception as e:
                resultados['fallidos'] += 1
                resultados['errores'].append(f"Registro {i+1}: {str(e)}")
        
        return resultados